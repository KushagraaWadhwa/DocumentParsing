import os
import docx
from sentence_transformers import SentenceTransformer, util

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        return ' '.join(full_text)
    except Exception as e:
        print(f"Error reading DOCX file {file_path}: {str(e)}")
        return ""

def read_text_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print(f"Error reading text file {file_path}: {str(e)}")
        return ""

def semantic_similarity(text1, text2):
    model = SentenceTransformer('all-MiniLM-L6-v2')
    embedding1 = model.encode(text1, convert_to_tensor=True)
    embedding2 = model.encode(text2, convert_to_tensor=True)
    return util.pytorch_cos_sim(embedding1, embedding2).item()

# directories = ["PyPDF2", "py_pdf_parser", "llama_parser","pymupdf"]
directories = ["ground_truth","pymupdf"]
files = ['policy1', 'policy2', 'handbook']

# Iterate over all pairs of directories and files
for dir1 in directories:
    for file1 in files:
        file_path1 = f'/Users/kushagrawadhwa/Desktop/PDF parser/{dir1}/{file1}_parsed.txt'
        text1 = read_text_file(file_path1)
        
        if not text1:
            continue
        
        for dir2 in directories:
            if dir1 == dir2:
                continue  # Skip comparing within the same directory
            
            for file2 in files:
                file_path2 = f'/Users/kushagrawadhwa/Desktop/PDF parser/{dir2}/{file2}_parsed.txt'
                text2 = read_text_file(file_path2)
                
                if not text2:
                    continue
                
                similarity_score = semantic_similarity(text1, text2)
                print(f"Similarity between {dir1}/{file1}_parsed.txt and {dir2}/{file2}_parsed.txt: {similarity_score:.4f}")

        print("\n")


# import os
# import docx
# from collections import Counter
# from sentence_transformers import SentenceTransformer, util

# # Function to extract unique words from docx file
# def extract_unique_words_from_docx(file_path):
#     try:
#         doc = docx.Document(file_path)
#         full_text = []
#         for para in doc.paragraphs:
#             full_text.append(para.text)
#         words = ' '.join(full_text).split()
#         unique_words = set(words)
#         return unique_words
#     except Exception as e:
#         return str(e)

# # Function to extract unique words from plain text file
# def extract_unique_words_from_txt(file_path):
#     try:
#         with open(file_path, 'r', encoding='utf-8') as file:
#             words = file.read().split()
#             unique_words = set(words)
#             return unique_words
#     except Exception as e:
#         return str(e)

# # Function to calculate semantic similarity
# def semantic_similarity(words1, words2):
#     model = SentenceTransformer('all-MiniLM-L6-v2')
#     embeddings1 = model.encode(list(words1), convert_to_tensor=True)
#     embeddings2 = model.encode(list(words2), convert_to_tensor=True)
#     cosine_scores = util.pytorch_cos_sim(embeddings1, embeddings2)
#     max_scores = cosine_scores.max(dim=1).values
#     return max_scores.mean().item()

# # List of document paths
# file_paths_docx = [f'/Users/kushagrawadhwa/Desktop/PDF parser/ground_truth/{file}.docx' for file in ("policy1", "policy2", "handbook")]
# file_paths_txt_dirs = ["PyPDF2", "py_pdf_parser", "llama_parser"]

# # Extract unique words and calculate similarity
# for file in ("policy1", "policy2", "handbook"):
#     file_path2 = f'/Users/kushagrawadhwa/Desktop/PDF parser/ground_truth/{file}.docx'
#     print(f"{file = }")
#     for dir in file_paths_txt_dirs:
#         file_path1 = f'/Users/kushagrawadhwa/Desktop/PDF parser/{dir}/{file}_parsed.txt'
#         print(f"{dir = }")

#         if os.path.exists(file_path2) and os.path.exists(file_path1):
#             docx_words = extract_unique_words_from_docx(file_path2)
#             txt_words = extract_unique_words_from_txt(file_path1)
            
#             if isinstance(docx_words, set) and isinstance(txt_words, set):
#                 similarity_score = semantic_similarity(docx_words, txt_words)
#                 print(f"Similarity score between the Word document and the parsed text file: {similarity_score:.2f}")
#             else:
#                 print(f"Error extracting unique words: {docx_words if isinstance(docx_words, str) else txt_words}")
#         else:
#             print(f"File not found: {file_path2 if not os.path.exists(file_path2) else file_path1}")
#         print("\n")